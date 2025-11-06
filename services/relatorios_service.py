from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from io import BytesIO
from datetime import datetime
import os
from sqlalchemy.orm import Session
from sqlalchemy import func, distinct
from services.graficos_service import AnalyticsService


class RelatorioService:
    @staticmethod
    def gerar_relatorio_geral_dashboard(db: Session, gestor_id: int) -> BytesIO:
        """
        Gera um relatório DOCX completo com todas as informações do dashboard do gestor.
        """
        try:
            # Buscar dados do dashboard
            dashboard_data = AnalyticsService.get_dashboard_data_for_gestor(db, gestor_id)
            gestor = dashboard_data['gestor']
            print(f"Dados do dashboard carregados com sucesso para gestor {gestor_id}")
            
            # Verificar se o gestor existe
            if gestor is None:
                raise ValueError(f"Gestor com ID {gestor_id} não encontrado no banco de dados")
                
        except Exception as e:
            print(f"Erro ao buscar dados do dashboard: {e}")
            raise
        
        # Criar documento
        doc = Document()
        
        # Configurar margens
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Adicionar logo (se existir)
        logo_path = "templates/static/img/DipePreta.svg"
        if os.path.exists(logo_path):
            try:
                # Para SVG, vamos usar uma abordagem diferente
                # Adicionar uma imagem placeholder ou texto do logo
                logo_paragraph = doc.add_paragraph()
                logo_run = logo_paragraph.add_run("DIPE")
                logo_run.font.size = Pt(24)
                logo_run.font.bold = True
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"Erro ao adicionar logo: {e}")
        
        # Título principal
        titulo = doc.add_heading('Relatório Geral', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Informações do relatório
        info_paragraph = doc.add_paragraph()
        info_paragraph.add_run(f"Gerado em: {datetime.now().strftime('%d/%m/%Y às %H:%M')}\n")
        info_paragraph.add_run(f"Gestor: {gestor.nome}\n")
        info_paragraph.add_run(f"Sistema: DIPE - Dashboard de Avaliações")
        info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adicionar linha separadora
        doc.add_paragraph("_" * 80)
        
        # 1. RESUMO EXECUTIVO
        doc.add_heading('1. Resumo Executivo', level=1)
        
        cards = dashboard_data['cards']
        resumo_table = doc.add_table(rows=1, cols=2)
        resumo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Cabeçalho da tabela
        hdr_cells = resumo_table.rows[0].cells
        hdr_cells[0].text = 'Indicador'
        hdr_cells[1].text = 'Valor'
        
        # Estilizar cabeçalho
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adicionar dados dos cards
        indicadores = [
            ('Total de Alunos', str(cards['total_alunos'])),
            ('Provas Realizadas', str(cards['total_provas'])),
            ('Média Geral', f"{cards['media_geral']} acertos"),
            ('Percentual Suficiente', f"{cards['percentual_suficiente']}%")
        ]
        
        for indicador, valor in indicadores:
            row_cells = resumo_table.add_row().cells
            row_cells[0].text = indicador
            row_cells[1].text = valor
            row_cells[0].paragraphs[0].runs[0].font.bold = True
        
        # 2. ANÁLISE POR DISCIPLINA
        doc.add_heading('2. Análise por Disciplina', level=1)
        
        try:
            desempenho = dashboard_data['desempenho_disciplina']
            print(f"Dados de desempenho: {desempenho}")
        except Exception as e:
            print(f"Erro ao acessar dados de desempenho: {e}")
            desempenho = {'labels': [], 'data': []}
        disciplina_table = doc.add_table(rows=1, cols=2)
        disciplina_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Cabeçalho
        hdr_cells = disciplina_table.rows[0].cells
        hdr_cells[0].text = 'Disciplina'
        hdr_cells[1].text = 'Média de Acertos'
        
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Dados das disciplinas
        try:
            for i, disciplina in enumerate(desempenho['labels']):
                row_cells = disciplina_table.add_row().cells
                row_cells[0].text = disciplina
                if i < len(desempenho['data']):
                    row_cells[1].text = f"{desempenho['data'][i]:.2f}"
                else:
                    row_cells[1].text = "0.00"
        except Exception as e:
            print(f"Erro ao processar dados de disciplina: {e}")
            # Adicionar linha de erro
            row_cells = disciplina_table.add_row().cells
            row_cells[0].text = "Erro ao carregar dados"
            row_cells[1].text = "N/A"
        
        # 3. DISTRIBUIÇÃO GEOGRÁFICA
        doc.add_heading('3. Distribuição Geográfica', level=1)
        
        # Por zona
        doc.add_heading('3.1 Por Zona', level=2)
        try:
            zona_data = dashboard_data['zona_distribuicao']
            print(f"Dados de zona: {zona_data}")
        except Exception as e:
            print(f"Erro ao acessar dados de zona: {e}")
            zona_data = {'labels': [], 'data': []}
        zona_table = doc.add_table(rows=1, cols=2)
        zona_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells = zona_table.rows[0].cells
        hdr_cells[0].text = 'Zona'
        hdr_cells[1].text = 'Quantidade de Alunos'
        
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for i, zona in enumerate(zona_data['labels']):
            row_cells = zona_table.add_row().cells
            row_cells[0].text = zona.title()
            if i < len(zona_data['data']):
                row_cells[1].text = str(zona_data['data'][i])
            else:
                row_cells[1].text = "0"
        
        # Por município
        doc.add_heading('3.2 Por Município', level=2)
        try:
            cidade_data = dashboard_data['cidade_distribuicao']
            print(f"Dados de cidade: {cidade_data}")
        except Exception as e:
            print(f"Erro ao acessar dados de cidade: {e}")
            cidade_data = {'labels': [], 'data': []}
        if cidade_data['labels']:
            cidade_table = doc.add_table(rows=1, cols=2)
            cidade_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            hdr_cells = cidade_table.rows[0].cells
            hdr_cells[0].text = 'Município'
            hdr_cells[1].text = 'Quantidade de Alunos'
            
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for i, cidade in enumerate(cidade_data['labels']):
                row_cells = cidade_table.add_row().cells
                row_cells[0].text = cidade
                if i < len(cidade_data['data']):
                    row_cells[1].text = str(cidade_data['data'][i])
                else:
                    row_cells[1].text = "0"
        else:
            doc.add_paragraph("Nenhum dado de município disponível.")
        
        # 4. PERFIL DOS ALUNOS
        doc.add_heading('4. Perfil dos Alunos', level=1)
        
        # Por idade
        doc.add_heading('4.1 Distribuição por Idade', level=2)
        idade_data = dashboard_data['perfil_idade']
        if idade_data['labels']:
            idade_table = doc.add_table(rows=1, cols=2)
            idade_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            hdr_cells = idade_table.rows[0].cells
            hdr_cells[0].text = 'Idade'
            hdr_cells[1].text = 'Quantidade de Alunos'
            
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for i, idade in enumerate(idade_data['labels']):
                row_cells = idade_table.add_row().cells
                row_cells[0].text = f"{idade} anos"
                if i < len(idade_data['data']):
                    row_cells[1].text = str(idade_data['data'][i])
                else:
                    row_cells[1].text = "0"
        else:
            doc.add_paragraph("Nenhum dado de idade disponível.")
        
        # 5. ANÁLISE DE DESEMPENHO
        doc.add_heading('5. Análise de Desempenho', level=1)
        
        # Distribuição das notas
        doc.add_heading('5.1 Distribuição das Notas', level=2)
        distribuicao = dashboard_data['distribuicao_notas']
        distribuicao_table = doc.add_table(rows=1, cols=2)
        distribuicao_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells = distribuicao_table.rows[0].cells
        hdr_cells[0].text = 'Classificação'
        hdr_cells[1].text = 'Quantidade'
        
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for i, classificacao in enumerate(distribuicao['labels']):
            row_cells = distribuicao_table.add_row().cells
            row_cells[0].text = classificacao
            if i < len(distribuicao['data']):
                row_cells[1].text = str(distribuicao['data'][i])
            else:
                row_cells[1].text = "0"
        
        # Taxa de participação
        doc.add_heading('5.2 Taxa de Participação', level=2)
        participacao = dashboard_data['taxa_participacao']
        participacao_table = doc.add_table(rows=1, cols=2)
        participacao_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells = participacao_table.rows[0].cells
        hdr_cells[0].text = 'Status'
        hdr_cells[1].text = 'Quantidade'
        
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for i, status in enumerate(participacao['labels']):
            row_cells = participacao_table.add_row().cells
            row_cells[0].text = status.title()
            if i < len(participacao['data']):
                row_cells[1].text = str(participacao['data'][i])
            else:
                row_cells[1].text = "0"
        
        # 6. ANÁLISE POR TURMAS
        doc.add_heading('6. Análise por Turmas', level=1)
        
        # Comparação entre turmas
        doc.add_heading('6.1 Comparação entre Turmas', level=2)
        turmas_data = dashboard_data['comparacao_turmas']
        if turmas_data['labels']:
            turmas_table = doc.add_table(rows=1, cols=2)
            turmas_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            hdr_cells = turmas_table.rows[0].cells
            hdr_cells[0].text = 'Turma/Curso'
            hdr_cells[1].text = 'Média de Acertos'
            
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for i, turma in enumerate(turmas_data['labels']):
                row_cells = turmas_table.add_row().cells
                row_cells[0].text = turma
                if i < len(turmas_data['data']):
                    row_cells[1].text = f"{turmas_data['data'][i]:.2f}"
                else:
                    row_cells[1].text = "0.00"
        else:
            doc.add_paragraph("Nenhum dado de turma disponível.")
        
        # Número de alunos por turma
        doc.add_heading('6.2 Número de Alunos por Turma', level=2)
        alunos_turma = dashboard_data['alunos_por_turma']
        if alunos_turma['labels']:
            alunos_table = doc.add_table(rows=1, cols=2)
            alunos_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            hdr_cells = alunos_table.rows[0].cells
            hdr_cells[0].text = 'Turma/Curso'
            hdr_cells[1].text = 'Quantidade de Alunos'
            
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for i, turma in enumerate(alunos_turma['labels']):
                row_cells = alunos_table.add_row().cells
                row_cells[0].text = turma
                if i < len(alunos_turma['data']):
                    row_cells[1].text = str(alunos_turma['data'][i])
                else:
                    row_cells[1].text = "0"
        
        # 7. TOP 10 ALUNOS
        doc.add_heading('7. Top 10 Alunos', level=1)
        top_alunos = dashboard_data['top_alunos']
        
        if top_alunos:
            top_table = doc.add_table(rows=1, cols=6)
            top_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Cabeçalho
            hdr_cells = top_table.rows[0].cells
            headers = ['Nome', 'Turma', 'Português', 'Matemática', 'Ciências', 'Média']
            
            for i, header in enumerate(headers):
                if i < len(hdr_cells):
                    hdr_cells[i].text = header
                    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                    hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Dados dos alunos
            for aluno in top_alunos:
                row_cells = top_table.add_row().cells
                row_cells[0].text = aluno['nome']
                row_cells[1].text = aluno['turma']
                row_cells[2].text = str(aluno['nota_portugues'])
                row_cells[3].text = str(aluno['nota_matematica'])
                row_cells[4].text = str(aluno['nota_ciencias'])
                row_cells[5].text = str(aluno['media'])
        else:
            doc.add_paragraph("Nenhum aluno com dados disponíveis.")
        
        # 8. PROGRESSÃO DOS ALUNOS
        doc.add_heading('8. Progressão dos Alunos', level=1)
        progressao = dashboard_data['progressao_alunos']
        
        if progressao['labels']:
            progressao_table = doc.add_table(rows=1, cols=2)
            progressao_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            hdr_cells = progressao_table.rows[0].cells
            hdr_cells[0].text = 'Ano'
            hdr_cells[1].text = 'Média de Acertos'
            
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for i, ano in enumerate(progressao['labels']):
                row_cells = progressao_table.add_row().cells
                row_cells[0].text = ano
                if i < len(progressao['data']):
                    row_cells[1].text = f"{progressao['data'][i]:.2f}"
                else:
                    row_cells[1].text = "0.00"
        else:
            doc.add_paragraph("Nenhum dado de progressão disponível.")
        
        # 9. CONCLUSÕES E RECOMENDAÇÕES
        doc.add_heading('9. Conclusões e Recomendações', level=1)
        
        conclusoes = doc.add_paragraph()
        conclusoes.add_run("Com base nos dados apresentados, observa-se que:\n\n")
        
        # Análise automática dos dados
        media_geral = float(cards['media_geral'])
        percentual_suficiente = float(cards['percentual_suficiente'])
        
        if media_geral >= 8:
            conclusoes.add_run("• O desempenho geral dos alunos está satisfatório, com média acima de 8 acertos.\n")
        elif media_geral >= 6:
            conclusoes.add_run("• O desempenho geral dos alunos está regular, com média entre 6 e 8 acertos.\n")
        else:
            conclusoes.add_run("• O desempenho geral dos alunos precisa de atenção, com média abaixo de 6 acertos.\n")
        
        if percentual_suficiente >= 70:
            conclusoes.add_run("• A taxa de alunos com desempenho suficiente está adequada (acima de 70%).\n")
        elif percentual_suficiente >= 50:
            conclusoes.add_run("• A taxa de alunos com desempenho suficiente está moderada (entre 50% e 70%).\n")
        else:
            conclusoes.add_run("• A taxa de alunos com desempenho suficiente está baixa (abaixo de 50%).\n")
        
        conclusoes.add_run("\nRecomendações:\n")
        conclusoes.add_run("• Analisar as disciplinas com menor desempenho para identificar necessidades de reforço.\n")
        conclusoes.add_run("• Implementar estratégias de apoio para alunos com dificuldades.\n")
        conclusoes.add_run("• Acompanhar regularmente a progressão dos alunos por turma.\n")
        conclusoes.add_run("• Incentivar a participação dos alunos nas avaliações.\n")
        
        # Rodapé
        doc.add_paragraph("\n" + "_" * 80)
        rodape = doc.add_paragraph()
        rodape.add_run(f"Relatório gerado automaticamente pelo sistema DIPE em {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
        rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Salvar em BytesIO
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        
        return buf

    @staticmethod
    def get_gestor_report_data(db: Session):
        """
        Método para compatibilidade com controllers existentes.
        Retorna dados básicos para relatórios do gestor.
        """
        # Buscar dados básicos
        from models.aluno import Aluno
        from models.prova import Prova
        from models.resultado import Resultado
        from sqlalchemy import func
        
        total_alunos = db.query(func.count(Aluno.idAluno)).scalar() or 0
        total_provas = db.query(func.count(distinct(Resultado.prova_id))).scalar() or 0
        
        # Calcular média geral
        resultados_acertos = db.query(Resultado.acertos).all()
        media_geral = 0
        if resultados_acertos:
            media_geral = sum(r.acertos for r in resultados_acertos) / len(resultados_acertos)

        return {
            'total_alunos': total_alunos,
            'total_provas': total_provas,
            'media_geral': f"{media_geral:.1f}",
            'data_geracao': datetime.now().strftime('%d/%m/%Y às %H:%M')
        }

    @staticmethod
    def get_professor_report_data(db: Session, professor_id: int):
        """
        Método para compatibilidade com controllers existentes.
        Retorna dados básicos para relatórios do professor.
        """
        from models.prova import Prova
        from models.resultado import Resultado
        from models.turma import Turma
        from dao.turma_dao import TurmaDAO
        from dao.banco_questoes_dao import BancoQuestoesDAO
        from sqlalchemy import func
        
        # Buscar estatísticas do professor
        turmas_count = len(TurmaDAO.get_active_by_professor(db, professor_id))
        questoes_count = BancoQuestoesDAO.get_count_by_professor(db, professor_id)
        provas_count = db.query(func.count(Prova.id)).filter(Prova.professor_id == professor_id).scalar()
        
        return {
            'turmas_count': turmas_count,
            'questoes_count': questoes_count,
            'provas_count': provas_count,
            'data_geracao': datetime.now().strftime('%d/%m/%Y às %H:%M')
        }
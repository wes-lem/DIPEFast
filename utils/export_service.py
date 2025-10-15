from typing import Dict, Any
from fastapi import Response

try:
	from weasyprint import HTML
	has_weasyprint = True
except Exception:
	has_weasyprint = False

from docx import Document
from docx.shared import Pt


def pdf_response_from_html(html_content: str, filename: str = "relatorio.pdf") -> Response:
	if not has_weasyprint:
		return Response(content=b"WeasyPrint nao instalado.", media_type="text/plain", status_code=500)
	pdf_bytes = HTML(string=html_content).write_pdf()
	headers = {"Content-Disposition": f"attachment; filename={filename}"}
	return Response(content=pdf_bytes, media_type="application/pdf", headers=headers)


def docx_response_from_data(title: str, sections: Dict[str, Any], filename: str = "relatorio.docx") -> Response:
	doc = Document()
	doc.add_heading(title, level=1)
	for section_title, section_data in sections.items():
		doc.add_heading(section_title, level=2)
		if isinstance(section_data, dict) and "labels" in section_data and ("medias" in section_data or "data" in section_data):
			labels = section_data.get("labels", [])
			values = section_data.get("medias", section_data.get("data", []))
			table = doc.add_table(rows=1, cols=2)
			hdr_cells = table.rows[0].cells
			hdr_cells[0].text = "Item"
			hdr_cells[1].text = "Valor"
			for label, value in zip(labels, values):
				row_cells = table.add_row().cells
				row_cells[0].text = str(label)
				row_cells[1].text = f"{float(value):.2f}"
		else:
			p = doc.add_paragraph(str(section_data))
			p.runs[0].font.size = Pt(11)
	
	from io import BytesIO
	buf = BytesIO()
	doc.save(buf)
	buf.seek(0)
	headers = {"Content-Disposition": f"attachment: filename={filename}"}
	return Response(content=buf.read(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)


